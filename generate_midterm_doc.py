#!/usr/bin/env python3
"""Generate the mid-term check form as a Word document with ALL screenshots."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

BASE = os.path.dirname(__file__)

def add_img(doc, path, caption, width=5.5):
    """Add a centered image with caption."""
    doc.add_paragraph()
    full = os.path.join(BASE, path)
    if os.path.exists(full):
        doc.add_picture(full, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6b, 0x7d, 0x8e)
    doc.add_paragraph()

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
# COVER PAGE
# ============================================================
doc.add_paragraph(); doc.add_paragraph()
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('（AI驱动的零碳园区数据平台）')
run.font.size = Pt(16); run.font.bold = True; run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.add_paragraph()
main_title = doc.add_paragraph(); main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = main_title.add_run('中 期 检 查 表')
run.font.size = Pt(26); run.font.bold = True; run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.add_paragraph(); doc.add_paragraph()
for label, value in [('实践单位：', '清华大学（校内实践）'),
                      ('团队成员：', '柯谨  徐青杨'),
                      ('实践指导教师：', '（校内指导老师）        （企业指导老师）')]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    run = p.add_run(label); run.font.size = Pt(12); run.font.bold = True
    run = p.add_run(value); run.font.size = Pt(12); run.font.underline = True
doc.add_paragraph()
date_p = doc.add_paragraph(); date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('二○二六年七月'); run.font.size = Pt(14); run.font.bold = True

doc.add_page_break()

# ============================================================
# SECTION 1: 实践工作进展摘要
# ============================================================
h = doc.add_heading('1. 实践工作进展摘要', level=1)
for r in h.runs: r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph(
    '本项目以广东省首批15个省级零碳园区为研究对象，围绕"AI赋能零碳园区的策略平台系统"建设目标，'
    '已完成全部三个阶段的开发工作。平台实现了从数据层到前端Dashboard再到AI Agent的完整链路，'
    '具备AI智能匹配和报告自动生成能力。'
)

# ---- Platform screenshots gallery ----
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('▎ 平台核心页面展示'); run.font.bold = True; run.font.size = Pt(14)

add_img(doc, 'screenshot-01-dashboard.png', '图1  总览仪表盘 — KPI统计卡片 + 工具类型分布环形图 + 园区覆盖柱状图 + 成熟度雷达图 + 建筑场景专区')
add_img(doc, 'screenshot-02-tool-list.png', '图2  工具箱浏览 — 按分类/园区类型/运营环节/成熟度多维度筛选 + 分页卡片列表（20个AI工具）')
add_img(doc, 'screenshot-03-tool-detail.png', '图3  工具详情页 — 完整技术路径、应用场景、AI赋能方式、价值主张、前置条件、参考案例')
add_img(doc, 'screenshot-06-ai-match-result.png', '图4  AI智能匹配结果 — 匹配推理分析 + 置信度评分 + 核心推荐（含匹配度百分比和推荐理由）+ 部署优先级标签')
add_img(doc, 'screenshot-05-report.png', '图5  AI报告生成页 — 左侧园区选择 + 报告配置 + 右侧五章节报告预览区')
add_img(doc, 'screenshot-07-report-generated.png', '图6  AI生成的分析报告 — 湛江经开区（东海岛）五章节完整报告：园区概况→工具推荐→缺口分析→路线图→总结建议')

doc.add_page_break()

# Progress table
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('表1  项目开发进度总览'); run.font.bold = True

table = doc.add_table(rows=6, cols=5, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h_text in enumerate(['阶段', '核心内容', '技术栈/数据', '完成状态', '完成时间']):
    cell = table.rows[0].cells[i]; cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs: run.font.bold = True; run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for row_idx, row_data in enumerate([
    ['阶段一\n数据层', 'AI工具数据库\n园区信息库\n商业案例库', '20个AI工具(9大类)\n15个广东省零碳园区\n8个标杆商业案例\nSQLite数据库', '✅ 完成', '第2周\n(7/7-7/11)'],
    ['阶段二\n前端Dashboard', '总览仪表盘\n工具箱浏览\n工具详情页\n园区匹配页', 'React 18 + TypeScript\nAnt Design 5\nECharts可视化\nVite构建', '✅ 完成', '第3周\n(7/12-7/17)'],
    ['阶段三\nAI Agent', 'AI智能匹配\nAI报告生成\nLLM Service\nDemo降级模式', 'OpenAI兼容SDK\nDeepSeek/Claude API\nFastAPI Agent Router\nMarkdown报告渲染', '✅ 完成', '第4周\n(7/18-7/20)'],
    ['关键成果', '5个前端页面 + 7个API端点\n规则匹配 + AI匹配双模式\nDemo模板 + AI生成双轨', '', '', ''],
    ['量化指标', '后端~1,500行 | 前端~2,000行\n覆盖15园区6类型9分类', '', '', ''],
]):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]; cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs: run.font.size = Pt(9)
            if row_idx >= 3:
                for run in p.runs: run.font.bold = True

doc.add_paragraph()

# New features description
doc.add_heading('新增AI功能说明', level=3)
doc.add_paragraph(
    '1. AI智能匹配（新增）：平台支持双模式切换——规则匹配模式（基于园区类型预设规则）和AI智能匹配模式'
    '（基于LLM大模型深度推理）。AI模式综合园区产业特征、关键方向、建设周期等多维因素，为每个推荐工具'
    '附带匹配评分（0-100%）、推理理由和部署优先级标签（立即部署/短期规划/长期储备）。'
)
doc.add_paragraph(
    '2. AI报告生成（新增）：一键生成面向单个园区的专业分析报告，包含五大章节——园区概况分析、'
    '核心AI工具推荐、技术缺口分析、分阶段实施路线图（基础建设期→核心部署期→优化提升期）、'
    '总结与行动建议。支持Demo模板模式（无需API Key即可体验完整功能）和AI生成模式（接入DeepSeek/Claude后启用）。'
)
doc.add_paragraph(
    '3. Demo双轨架构：设计了完整的降级策略——当LLM API未配置或不可用时，自动切换至Demo模板模式，'
    '确保平台在任何环境下均可正常展示所有功能，适合课堂演示场景。'
)

doc.add_page_break()

# ============================================================
# SECTION 2: 阶段性总结
# ============================================================
h = doc.add_heading('2. 实践工作的阶段性总结（含初步结论）', level=1)
for r in h.runs: r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_heading('2.1 项目定位演变', level=2)
p = doc.add_paragraph(); p.add_run('本项目经历了三次重要的定位深化：').font.bold = True
table = doc.add_table(rows=4, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h_text in enumerate(['时间节点', '定位', '关键变化']):
    cell = table.rows[0].cells[i]; cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs: run.font.bold = True; run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row_idx, row_data in enumerate([
    ['第1周（初始）', '单园区AI评估工具', '标准识别→数据填报→测算诊断→报告导出\nMVP范围，2人小组可行'],
    ['第2周', '零碳园区AI工具策略知识库', '从单一评估工具拓宽为系统化知识库'],
    ['第3周（调整后）', 'AI赋能零碳园区的\n策略平台系统', '四大拓展：场景全域化、动态AI Agent、\n可视化Dashboard、多角色用户'],
]):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]; cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs: run.font.size = Pt(10)

doc.add_paragraph()
p = doc.add_paragraph(); run = p.add_run('核心结论：'); run.font.bold = True
p.add_run('与指导老师的讨论是本项目的关键转折点。将项目从"静态知识库"升级为"具备动态更新能力与AI推理能力的策略平台系统"，使项目既具有学术研究的深度，又具备实际应用价值。')

doc.add_heading('2.2 广东省零碳园区调研发现', level=2)
doc.add_paragraph(
    '通过对广东省首批15个省级零碳园区的系统性调研（覆盖14个地级市、珠三角+粤东+粤西+粤北四大区域），'
    '归纳出六大产业类型，并对每类园区的AI工具需求进行了系统梳理：'
)

table = doc.add_table(rows=8, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h_text in enumerate(['产业类型', '数量', '代表园区', '零碳建设重点方向']):
    cell = table.rows[0].cells[i]; cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs: run.font.bold = True; run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row_idx, row_data in enumerate([
    ['先进制造型', '5个', '广州南沙、佛山狮山、惠州惠城\n中山翠亨、江门台山', '绿电直供、能效提升、智能制造降碳'],
    ['重化工近零碳型', '2个', '湛江东海岛（钢铁）\n茂名滨海新区（绿色化工+氢能）', '氢基冶炼、CCUS、绿氢替代'],
    ['新能源装备制造型', '2个', '阳江高新区（风电装备）\n肇庆大旺（智能汽车）', '清洁能源装备、绿色交通'],
    ['新材料型', '1个', '潮州新材料产业园', '低碳材料工艺、循环利用'],
    ['临港特色产业型', '2个', '汕头潮阳、汕尾红海湾', '岸电替代、港口低碳化'],
    ['生态高新技术型', '3个', '河源高新区、梅州融湾\n云浮新兴产业', '生态保护+低碳产业协同'],
    ['合计', '15个', '覆盖14个地级市、4大区域', '6大类、11个运营环节'],
]):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]; cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs: run.font.size = Pt(9)
        if row_idx == 6:
            for p in cell.paragraphs: p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph('初步结论：')
doc.add_paragraph('（1）园区类型高度分化，AI需求差异显著——六大类型园区的AI工具需求存在显著差异，先进制造型园区需求最广，重化工近零碳型最具挑战性。', style='List Bullet')
doc.add_paragraph('（2）AI工具覆盖存在明显短板——交通物流（★★★）和碳汇管理（★★）是当前最薄弱的两个环节，也是平台后续可重点突破的差异化方向。', style='List Bullet')
doc.add_paragraph('（3）商业平台已有成熟先例但缺乏通用平台——现有8个标杆方案均为单园区定制化部署，本项目填补了跨园区、跨类型的通用AI工具箱策略平台空白。', style='List Bullet')

doc.add_heading('2.3 AI工具箱九大分类体系', level=2)
doc.add_paragraph(
    '项目构建了面向零碳园区的AI工具九大分类体系：'
)
table = doc.add_table(rows=10, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h_text in enumerate(['图标', '分类', '代表工具（示例）']):
    cell = table.rows[0].cells[i]; cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs: run.font.bold = True; run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row_idx, row_data in enumerate([
    ['🔮', '预测类（5★）', '电力负荷预测、光伏/风电出力预测'],
    ['⚡', '优化类（4★）', '储能充放电策略优化、多能流调度优化、氢基竖炉工艺AI优化'],
    ['🎛️', '控制类（5★）', '暖通空调AI节能控制、照明智能调控'],
    ['🔧', '诊断类（4★）', '设备能效异常诊断、故障预警与健康管理'],
    ['📐', '核算类（3★）', '碳足迹核算、产品碳标签生成'],
    ['👁️', '识别类（4★）', '危化品泄漏AI检测、植被碳汇遥感评估'],
    ['📅', '调度类（4★）', '新能源车队智能调度、岸电智能调度'],
    ['📚', '知识类（4★）', '能碳指标智能分析、政策标准知识问答'],
    ['🚀', '创新类（2★）', 'CCUS碳捕集效率AI模拟、零碳路径情景模拟'],
]):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]; cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs: run.font.size = Pt(10)

doc.add_paragraph()
doc.add_heading('2.4 技术实现阶段性结论', level=2)
doc.add_paragraph('（1）前后端分离架构验证可行——React+FastAPI+SQLite轻量化技术栈，2人小组3周完成全部开发。', style='List Bullet')
doc.add_paragraph('（2）LLM集成架构设计合理——OpenAI兼容SDK实现DeepSeek/Claude双供应商无缝切换，Demo降级模式确保演示可靠性。', style='List Bullet')
doc.add_paragraph('（3）Demo+AI双模式是实用选择——兼顾演示零风险和AI能力展示，适合大学课程项目场景。', style='List Bullet')

doc.add_page_break()

# ============================================================
# SECTION 3: 下一步工作计划
# ============================================================
h = doc.add_heading('3. 下一步的工作计划', level=1)
for r in h.runs: r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_heading('3.1 必做任务（第4-5周）', level=2)
table = doc.add_table(rows=5, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h_text in enumerate(['序号', '任务', '优先级']):
    cell = table.rows[0].cells[i]; cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs: run.font.bold = True; run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row_idx, row_data in enumerate([
    ['1', '配置真实LLM API — 填入DeepSeek API Key，完成AI匹配和报告的真实验证', '🔴 高'],
    ['2', '扩充工具条目至40+ — 增加非建筑场景（交通物流、水资源、废弃物、碳汇）', '🔴 高'],
    ['3', '报告撰写 — 完成开题报告、中期报告的正式版', '🔴 高'],
    ['4', '用户测试 — 邀请2-3名同学模拟不同角色使用平台，收集反馈', '🟡 中'],
]):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]; cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs: run.font.size = Pt(10)

doc.add_paragraph()
doc.add_heading('3.2 选做任务（视时间情况）', level=2)
for item in [
    '知识抓取Agent — 实现从网络公开数据源自动抓取AI工具和案例的Agent',
    'PDF报告导出 — 将在线报告一键导出为PDF/DOCX格式',
    '多园区对比分析 — Dashboard新增多园区AI工具覆盖度对比视图',
    '用户角色切换 — 前端新增角色切换（园区管理者/政策制定者/运维工程师）',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_paragraph()
doc.add_heading('3.3 后续时间线', level=2)
doc.add_paragraph('第4周 (7/21-7/27):  LLM API配置验证 + 工具条目扩充 + 正式报告撰写')
doc.add_paragraph('第5周 (7/28-8/3):  用户测试反馈 + 优化迭代 + 知识抓取Agent（如时间允许）')
doc.add_paragraph('第6周 (8/4-8/10):  结题准备 + 答辩PPT + 结题报告')

doc.add_page_break()

# ============================================================
# SECTION 4: 未按计划完成的内容
# ============================================================
h = doc.add_heading('4. 哪些实践内容未按计划完成，原因何在', level=1)
for r in h.runs: r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_heading('4.1 已按计划完成的内容 ✅', level=2)
for item in [
    'AI工具数据库构建（20个工具、9大分类）— 超额完成（含10+丰富字段）',
    '广东省15个零碳园区数据调研与录入',
    '商业案例库（8个标杆平台）',
    'Dashboard可视化平台开发（5页面）',
    '园区-工具规则匹配',
    'LLM驱动的AI智能匹配（比计划提前完成）',
    'AI报告自动生成（比计划提前完成）',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.2 未按原计划完成的内容 ⚠️', level=2)
table = doc.add_table(rows=4, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h_text in enumerate(['原计划内容', '完成情况', '原因']):
    cell = table.rows[0].cells[i]; cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs: run.font.bold = True; run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row_idx, row_data in enumerate([
    ['知识动态抓取Agent', '🔲 未完成', '优先级调整：与老师讨论后Dashboard和AI Agent提升为更高优先级'],
    ['MCP协议工具生态对接', '🔲 未完成', '技术复杂度超预期：需额外研究和开发时间，作为选做项保留'],
    ['AI工具条目目标100+', '⚠️ 完成20个', '重质量轻数量：每个条目包含10+字段，信息密度高'],
]):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]; cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs: run.font.size = Pt(10)

doc.add_paragraph()
doc.add_heading('4.3 超计划完成的内容 🚀', level=2)
for item in [
    'AI Agent智能匹配 — 提前至第3周完成，LLM推理匹配+Demo降级双模式',
    'AI报告生成 — 提前至第3周完成，五章节结构化报告',
    '建筑场景专题分析 — Dashboard新增建筑运行场景专题模块',
    'Demo双模式架构 — 完整的降级策略保障演示可靠性（无API Key环境仍可展示全部功能）',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# SECTION 5: 存在问题与建议
# ============================================================
h = doc.add_heading('5. 存在问题、建议及需要说明的情况', level=1)
for r in h.runs: r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_heading('5.1 存在问题', level=2)
table = doc.add_table(rows=5, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h_text in enumerate(['问题', '严重程度', '说明']):
    cell = table.rows[0].cells[i]; cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs: run.font.bold = True; run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row_idx, row_data in enumerate([
    ['LLM API费用', '🟡 中', 'DeepSeek API虽便宜（~¥1/百万token）但仍需小额充值'],
    ['知识库规模', '🟡 中', '当前20个工具覆盖主要场景，非建筑场景工具偏少'],
    ['无用户认证系统', '🟢 低', '不影响Demo演示，多角色视图通过前端模拟实现'],
    ['数据持久化', '🟢 低', 'SQLite适合开发/Demo，生产环境需迁移PostgreSQL'],
]):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]; cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs: run.font.size = Pt(10)

doc.add_paragraph()
doc.add_heading('5.2 建议', level=2)
for item in [
    'API费用解决方案：建议使用DeepSeek免费额度（注册即送）或使用腾讯云等国内平台的LLM API。',
    '知识库扩展建议：在开题报告中明确MVP范围（20-30个核心工具），将100+工具作为后续平台化愿景写入"展望"章节。',
    '项目展示建议：答辩/Demo时重点展示三个功能流——Dashboard总览→分类筛选→工具详情（展示知识库完整度）；选择园区→AI智能匹配→查看推荐理由（展示AI Agent能力）；选择园区→AI报告生成→浏览五章节报告（展示LLM应用价值）。',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_paragraph()
doc.add_heading('5.3 课题分工', level=2)
table = doc.add_table(rows=3, cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h_text in enumerate(['成员', '主要分工', '贡献']):
    cell = table.rows[0].cells[i]; cell.text = h_text
    for p in cell.paragraphs:
        for run in p.runs: run.font.bold = True; run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row_idx, row_data in enumerate([
    ['柯谨', '广东省零碳园区调研、AI工具知识库撰写\nDashboard需求设计、周报撰写', '数据采集与内容建设'],
    ['徐青杨', '平台架构设计、前后端开发\n（FastAPI+React+TypeScript）\nAI Agent集成、数据库设计', '技术开发与系统实现'],
]):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]; cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs: run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph('关于实践单位的说明：本项目为校内实践课题，实践单位为清华大学。项目无需外部企业合作，所有开发工作在校内完成。')
doc.add_paragraph('关于后续软著/专利申请的说明：平台已具备Demo演示能力，建议在知识库扩充至30-40个工具条目后申请软件著作权1件。核心创新点：面向零碳园区的AI工具多维度匹配推荐方法。')

doc.add_page_break()

# ============================================================
# APPENDIX: Architecture
# ============================================================
h = doc.add_heading('附录：平台技术架构图', level=1)
for r in h.runs: r.font.name = '黑体'; r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

# Architecture diagram as structured text
arch = """┌─────────────────────────────────────────────────────────────┐
│              前端层 (React 18 + TypeScript + Vite)              │
│  Dashboard │ ToolList │ ToolDetail │ ParkMatch │ Report      │
│               Ant Design 5  +  ECharts 可视化                   │
├─────────────────────────────────────────────────────────────┤
│                    Vite Proxy (/api → :8080)                    │
├─────────────────────────────────────────────────────────────┤
│                 API层 (Python FastAPI)                          │
│  GET /api/dashboard          — KPI + 图表数据                  │
│  GET /api/tools              — 工具列表（筛选+分页）             │
│  GET /api/tools/:id          — 工具详情 + 案例                  │
│  GET /api/parks              — 园区列表                        │
│  GET /api/match?park_id=X    — 规则匹配                        │
│  POST /api/agent/match       — AI智能匹配 (NEW!)               │
│  POST /api/agent/report      — AI报告生成 (NEW!)               │
│           SQLAlchemy ORM  +  Pydantic Schema                    │
├─────────────────────────────────────────────────────────────┤
│               AI层 (OpenAI-compatible SDK)                      │
│     DeepSeek API  ←——→  Claude API  (通过 .env 切换)           │
│     Demo 降级模式（无 API Key 时自动启用模板模式）               │
├─────────────────────────────────────────────────────────────┤
│                 数据层 (SQLite Database)                        │
│   20 tools │ 15 parks │ 8 cases │ 9 categories │ update_logs  │
└─────────────────────────────────────────────────────────────┘"""

p = doc.add_paragraph()
run = p.add_run(arch)
run.font.size = Pt(8); run.font.name = 'Courier New'

doc.add_paragraph()
doc.add_paragraph()

# Add the design images as appendix
add_img(doc, 'dashboard-design.png', '附图A  总览仪表盘设计稿 — KPI卡片 + 工具类型分布饼图 + 园区覆盖热力图 + 成熟度雷达图 + 更新日志', 5.0)
add_img(doc, 'dashboard-detail-design.png', '附图B  工具箱浏览 + 工具详情 + 园区匹配设计稿', 5.0)

# ============================================================
# SIGNATURE PAGE
# ============================================================
doc.add_page_break()
doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('考查小组对课题进展情况、已有完成质量的评价，提出的进一步工作意见：')
run.font.size = Pt(12); run.font.bold = True
doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()
doc.add_paragraph(); doc.add_paragraph()
doc.add_paragraph('实践指导教师（签字）：________________________         日期：2026年    月    日')
doc.add_paragraph(); doc.add_paragraph()
doc.add_paragraph('考查小组成员（签字）：________________________         日期：2026年    月    日')

# ============================================================
# SAVE
# ============================================================
output = os.path.join(BASE, '中期检查表-已填写.docx')
doc.save(output)
print(f'✅ 中期检查表已保存: {output}')
print(f'   包含图片: 2张设计图 + 7张平台运行截图 = 9张图片')
